import requests 
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor 
class Parser:
    source_link="https://en.wikipedia.org/wiki/"
    def __init__(self,project_topic):
        self.project_topic = project_topic
        self.project_paras=0
        self.completed=None
        self.collection_paragraphs=None#it will contain the whole data
        self.para_to_be_docxed=[]

    def parse(self):
        response=requests.get(Parser.source_link+self.project_topic)
        if response.status_code!=200:
            print("""Some problem occured... Plz make sure the content is heading is correct. If correct then some connection issue"""
)
            return """Some problem occured... Plz make sure the content is heading is correct. If correct then some connection issue"""
        soup=BeautifulSoup(response.content,"html.parser")
        body=soup.body
        #deleting unnecessary content
        try:
            for i in body.find_all(class_="reference"):
                i.decompose()
        except:
            pass
        
        # parsing html paragraphs to text
        parsed_pragraphs="" #to store parsed paragraphs
        number_of_para=0 #to count number of para
        para_list=[]#it will contain a list of all paragrahs stored in different tuples
        saving_list=[]
        html_para=body.find_all("p")
        
        for para in html_para:
            para_list_pointer=""
            number_of_para+=1
            parsed_pragraphs+=para.text
            para_list_pointer+=para.text
            saving_list.append(para_list_pointer)
            para_list_pointer=para_list_pointer.encode("utf-8","ignore")
            para_list.append((para_list_pointer))
            para_list_pointer=""

        self.project_paras=number_of_para
        self.completed=parsed_pragraphs
        self.collection_paragraphs=para_list
        self.para_to_be_docxed=saving_list

    @staticmethod
    def save_docx(file,collection_paragraphs,colors=[]):
        file=file.upper()
        document = Document()
        document.add_heading(file, 0)
        color_count=0
        if len(colors)==0:
            colors=["000000"]# #000000 should be 000000 and it means black
    
        for i in collection_paragraphs:
            if i.strip()=="":
                continue
            else:
                color_count+=1
                if color_count>(len(colors)-1):
                    color_count=0
                para=document.add_paragraph().add_run(str(i))
                para.font.color.rgb = RGBColor.from_string(colors[color_count])
                para.font.size=Pt(12)
        document.save(f'Your Projects/{file}.docx')
        
    @staticmethod
    def save_txt(self,file):
        with open(f"{file}.txt","w") as f:
            f.write(self.completed)
    def __repr__(self) :
        return(str(self.completed.encode("utf-8","ignore")))                

if __name__=="__main__":##to execute the file when it will be running as program not as a module 
    obj1=Parser("cricket")
    obj1.parse()
    # obj1.save("text.text")
    obj1.save_docx("cricket",obj1.para_to_be_docxed,colors=['0000a0', '000000', 'ff0080', '000000', '400080'])

